import sys
import os
from docx import Document

def read_docx(file_path):
    """
    Reads a .docx file and prints its content.
    """
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        sys.exit(1)

    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append('\t'.join(row_text))
        
        # Use sys.stdout.buffer.write to handle UTF-8 output more reliably on Windows
        output = '\n'.join(full_text)
        sys.stdout.buffer.write(output.encode('utf-8'))
        sys.stdout.buffer.write(b'\n')
    except Exception as e:
        print(f"Error reading file: {e}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <path_to_docx>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    read_docx(file_path)
